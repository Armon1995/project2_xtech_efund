"""
Created on Sat Mar 1 14:30:59 2024

Author: davideliu

E-mail: davide97ls@gmail.com

Goal: Generate Word report
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor  # 设置字体颜色
import re
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import pandas as pd
from io import StringIO
from datetime import datetime


# 可以将markdown里面的table加载word里面
def add_table(doc, markdown_table):
    """Adds a table to the document."""
    df = pd.read_csv(StringIO(markdown_table), sep="|").iloc[:, 1:-1]  # 去掉多余的空列
    df = df.drop(index=0).reset_index(drop=True)

    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    table.style = 'Table Grid'

    # 添加表头
    for j, col_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col_name.strip()
        run = cell.paragraphs[0].runs[0]
        run.font.name = u'楷体'  # 确保字体名称正确
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        run.font.size = Pt(12)
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加内容
    for i, row in df.iterrows():
        for j, cell_value in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(cell_value).strip()
            run = cell.paragraphs[0].runs[0]
            run.font.name = u'楷体'  # 确保字体名称正确
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
            run.font.size = Pt(10)
            if j == 0:
                run.font.bold = True
            if isinstance(cell_value, (int, float)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 设置边框函数
    def set_table_border(table):
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))  # 尝试获取已有的 tblPr 元素
        if tblPr is None:  # 如果没有，创建一个新的 tblPr 元素
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        borders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # 设置边框类型为单线
            border.set(qn('w:sz'), '4')        # 设置边框大小
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 设置边框颜色为黑色
            borders.append(border)

        tblPr.append(borders)

    set_table_border(table)
    table.autofit = True


def add_heading(document, text, level=1):
    """添加标题到文档中。
    Args:
        document: Document 对象。
        text: 标题文本。
        level: 标题级别，0-9。
    """
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = u'宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = Pt(20-level*2)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 设置字体颜色为黑色


def add_paragraph(document, text, indent=True):
    """添加正文段落到文档中。
    Args:
        document: Document 对象。
        text: 正文文本。
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = u'楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run.font.size = Pt(14)
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.3)


def add_picture(document, image_path, s=True):
    """添加图像到文档中并自动适应页面宽度。
    Args:
        document: Document 对象。
        image_path: 图像文件路径。
    """
    # 获取文档页面宽度（减去左右页边距）
    section = document.sections[-1]
    page_width = section.page_width - section.left_margin - section.right_margin

    # 获取图片的实际宽高
    image = Image.open(image_path)
    image_width, image_height = image.size

    # 计算缩放比例，使图片宽度适应页面宽度
    scale = page_width / image_width

    # 插入图片，按比例设置宽高
    if s:
        document.add_picture(image_path, width=page_width, height=image_height * scale)
    else:
        document.add_picture(image_path, width=page_width, height=image_height)

    # 设置图片居中
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# 将不同大小的标题数据取出来
def chunking_data(text):
    """Splits the data into chunks based on markdown headers."""
    matches = re.findall(r'# (.*?)\n(.*?)(?=\n# |\Z)', text, re.S)
    return matches


def chunking_data1(text):
    """Splits the data into chunks based on markdown headers."""
    matches = re.findall(r'## (.*?)\n(.*?)(?=\n## |\Z)', text, re.S)
    return matches


def generate_title(date_str):
    """Generates the title for the document based on the date."""
    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    year = date_obj.year
    month = date_obj.month
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    title = f"AI研报：{year}年{month}月预测{next_year}年{next_month}月LPR是否会下降"
    return title


def generate_word_doc(date: str):
    """
    Generates a Word document based on the provided date.

    This function reads various data files from a specified directory, processes the data,
    and creates a structured Word document with headings, paragraphs, images, and tables.

    Args:
        date (str): The date for which the report is being generated. This is used to locate the data files.

    Returns:
        None
    """
    doc = Document()

    # 添加标题
    title = generate_title(date)
    print('Title: ', title)
    add_heading(doc, title, level=0)

    # 添加引言
    with open(f'test_results/{date}/引言.md', 'r', encoding='utf-8') as f:
        background_data_analysis = f.read()
    background_data_analysis = chunking_data1(background_data_analysis)

    add_heading(doc, "一、引言", level=1)
    add_heading(doc, '1、背景介绍', level=2)
    add_paragraph(doc, background_data_analysis[0][1].strip().replace('\n\n', '\n'))
    add_heading(doc, '2、研究员的局限性', level=2)
    add_paragraph(doc, background_data_analysis[1][1].strip().replace('\n\n', '\n'))
    # print((background_data_analysis[1][1].strip().replace('\n\n','\n')))
    add_heading(doc, '3、人工智能分析的优势', level=2)
    add_paragraph(doc, background_data_analysis[2][1].strip().replace('\n\n', '\n'))

    # 添加LPR数据分析
    with open(f'test_results/{date}/LPR数据分析研报部分.md', 'r', encoding='utf-8') as f:
        y_data_analysis = f.read()
    y_data_analysis = chunking_data(y_data_analysis)

    add_heading(doc, "一、LPR介绍与分析", level=1)
    add_heading(doc, '1、LPR概述及其重要性', level=2)
    add_paragraph(doc, y_data_analysis[0][1])
    add_heading(doc, '2、LPR的变化趋势', level=2)
    add_picture(doc, f'test_results/{date}/LPR历史数据.png')
    add_paragraph(doc, y_data_analysis[1][1], False)
    add_heading(doc, '3、LPR趋势分析', level=2)
    add_paragraph(doc, y_data_analysis[2][1])
    add_heading(doc, '4、经济洞察与未来展望', level=2)
    add_paragraph(doc, y_data_analysis[3][1].replace('\n', ''))

    # 添加X数据分析
    with open(f'test_results/{date}/X数据分析研报部分.md', 'r', encoding='utf-8') as f:
        x_data_analysis = f.read()
    x_data_analysis = chunking_data(x_data_analysis)
    importance_factor_analysis = chunking_data1(x_data_analysis[1][1])
    add_heading(doc, "二、经济因素与LPR的关联分析", level=1)
    add_heading(doc, '1、相关经济因素对LPR的影响', level=2)
    add_paragraph(doc, x_data_analysis[0][1].replace('\n', ''))
    add_picture(doc, f'test_results/{date}/Xx_correlation_report.png')
    add_picture(doc, f'test_results/{date}/Xy_correlation_report.png')
    add_heading(doc, '2、主要因素分析', level=2)
    for i in range(len(importance_factor_analysis)):
        add_paragraph(doc, importance_factor_analysis[i][0]+'\n'+importance_factor_analysis[i][1].replace('**', '').
                      replace('  -', '  -').replace('：\n', '：').replace('  ', '').replace('\n\n', '\n'), False)
        add_picture(doc, f'test_results/{date}/top{i}相关因素分析.png')
    add_heading(doc, '3、其余因素分析', level=2)
    add_paragraph(doc, x_data_analysis[2][1].replace('\n', ''))

    add_heading(doc, '4、总结', level=2)
    add_paragraph(doc, x_data_analysis[3][1].replace('\n', ''))

    # 添加报告数据分析
    with open(f'test_results/{date}/报告对比分析研报部分.md', 'r', encoding='utf-8') as f:
        report_data_analysis = f.read()
    report_data_analysis = chunking_data(report_data_analysis)
    # assist_report_analysis = chunking_data1(report_data_analysis[0][1])
    add_heading(doc, "三、宏观政策和LPR的关联分析", level=1)
    add_paragraph(doc, report_data_analysis[0][1].replace('\n', '').replace('##', '\n'))
    add_heading(doc, '1、重点报告分析', level=2)
    importance_report_part = report_data_analysis[1][1]
    start = importance_report_part.find('##') + 2
    end = importance_report_part.find('|')
    substring = importance_report_part[start:end]
    add_heading(doc, substring, level=3)
    table_end = importance_report_part.rfind('|')
    tabel_text = importance_report_part[end:table_end+1]
    add_table(doc, tabel_text)
    importance_report_summary = importance_report_part[table_end+1:]
    add_paragraph(doc, importance_report_summary.replace('\n', ''))
    add_heading(doc, '文本分析', level=3)
    with open(f'test_results/{date}/report_wordcloud.txt', 'r', encoding='utf-8') as f:
        wordcloud_text = f.read()
    add_paragraph(doc, wordcloud_text)
    add_picture(doc, f'test_results/{date}/report_wordcloud.png')
    with open(f'test_results/{date}/terms_sentiment_bar_chart.md', 'r', encoding='utf-8') as f:
        terms_sentiment_bar_chart = f.read()
    add_paragraph(doc, terms_sentiment_bar_chart)
    add_picture(doc, f'test_results/{date}/terms_sentiment_bar_chart.png')

    add_heading(doc, '2、其余报告分析', level=2)
    add_paragraph(doc, report_data_analysis[2][1].strip().replace('\n\n', '\n'))
    add_heading(doc, '3、总结', level=2)
    add_paragraph(doc, report_data_analysis[3][1].replace('\n', ''))

    # 添加新闻数据分析
    with open(f'test_results/{date}/新闻数据分析研报部分.md', 'r', encoding='utf-8') as f:
        news_data_analysis = f.read()
    news_data_analysis = chunking_data(news_data_analysis)
    detail_news_analysis = chunking_data1(news_data_analysis[1][1])
    add_heading(doc, "四、新闻对LPR的影响分析", level=1)
    add_heading(doc, '1、新闻重要性', level=2)
    add_paragraph(doc, news_data_analysis[0][1].replace('\n', ''))
    add_heading(doc, '2、近期关键新闻分析', level=2)
    for i in range(len(detail_news_analysis)):
        add_paragraph(doc, detail_news_analysis[i][1].replace('\n\n', ':'))
    add_heading(doc, '3、总结分析结果', level=2)
    add_paragraph(doc, news_data_analysis[2][1])

    # 添加结果分析
    with open(f'test_results/{date}/reflection结果.md', 'r', encoding='utf-8') as f:
        result = f.read()
    result = chunking_data(result)
    detail_analysis = chunking_data1(result[1][1])
    add_heading(doc, '六、总体降息信号分析', level=1)
    add_paragraph(doc, result[0][1])
    for analysis in detail_analysis:
        add_heading(doc, analysis[0], level=2)
        add_paragraph(doc, analysis[1].replace('\n', '\n  '), False)

    # 添加附录
    add_heading(doc, '附录', level=1)
    add_heading(doc, '1、所有经济因素的重要性', level=2)
    add_picture(doc, f'test_results/{date}/各经济因素导致LPR下降的概率.png')  # removed due to image not generated
    add_heading(doc, '2、所有报告详细对比分析', level=2)
    with open(f'test_results/{date}/货币政策委员会会议分析.md', 'r', encoding='utf-8') as f:
        report_data_analysis = f.read()
    table_end = report_data_analysis.rfind('|')
    tabel_text = report_data_analysis[0:table_end+1]
    add_heading(doc, '货币政策委员会会议分析', level=3)
    add_table(doc, tabel_text)

    with open(f'test_results/{date}/货币政策分析.md', 'r', encoding='utf-8') as f:
        report_data_analysis = f.read()
    table_end = report_data_analysis.rfind('|')
    tabel_text = report_data_analysis[0:table_end+1]
    add_heading(doc, '货币政策执行报告分析', level=3)
    add_table(doc, tabel_text)

    with open(f'test_results/{date}/政治局会议分析.md', 'r', encoding='utf-8') as f:
        report_data_analysis = f.read()
    table_end = report_data_analysis.rfind('|')
    tabel_text = report_data_analysis[0:table_end+1]
    add_heading(doc, '政治局会议分析', level=3)
    add_table(doc, tabel_text)

    # 保存doc
    doc_path = f'test_results/{date}/{date}_report.docx'
    doc.save(doc_path)
    print(f'Doc generated and saved to: {doc_path}')


if __name__ == '__main__':
    dates = ['2025-02-28']
    for date in dates:
        print(f'Creating Word doc date {date}...')
        generate_word_doc(date)
        print(f"Word doc {date} completed.")
